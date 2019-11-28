from docx import Document
from docx.shared import Pt, Mm
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE
from docx.oxml.shared import OxmlElement, qn
from PIL import Image

Razd_simv = '-'



def Word(inf, lbsz, f, f2, fname, name_vl, NS, var,grop,zpps,per_name,D_file, Trig,tablepo):
    # Создаём объект документа
    if Trig: #'template_word/default.docx'
        document = D_file
    else:
        document = Document() 

    # Наследуем стиль и изменяем его
    style = document.styles['Normal'] # Берём стиль Нормальный

    f0 = style.font # Переменная для изменения параметров стиля
    f0.name = 'Arial' # Шрифт
    f0.size = Pt(12) # Размер шрифта

    pf = style.paragraph_format
    pf.line_spacing = Pt(0) # Междустрочный интервал
    pf.space_after = Pt(0) # Интервал после абзаца

    # Задаём свойство полей
    sections = document.sections
    s = sections[0]
    s.left_margin = Mm(30)
    s.right_margin = Mm(15)
    s.top_margin = Mm(20)
    s.bottom_margin = Mm(20)
    s.page_height = Mm(297)
    s.page_width = Mm(210)

    if tablepo =='Выводить таблицы и эпюры':
        tb1=True
        tb2=True
    elif tablepo =='Выводить таблицы':
        tb1=True
        tb2=False
    elif tablepo =='Выводить эпюры':
        tb1=False
        tb2=True

    
    # Функция выравнивания содержимого таблицы по вертикали
    def set_cell_vertical_alignment(cell, align="center"): 
        try:   
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcValign = OxmlElement('w:vAlign')  
            tcValign.set(qn('w:val'), align)  
            tcPr.append(tcValign)
            return True 
        except:
            #traceback.print_exc()             
            return False
    p0=['','','','']
    sv='.'+var if var!='' else ''
    if tb1:
    
        # Зоголовок 
        if not Trig:
            document._body.clear_content() # Удаляет все параграфы(Обезаетльно протом нужно чтото добавить)
        p0[0] = document.add_paragraph()
        p0[0].add_run('РЕЖИМ ЗАЗЕМЛЕНИЯ').bold = True
        pt_f1 = p0[0].paragraph_format  
        pt_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_f1.space_after = Pt(2)


        p0[1] = document.add_paragraph()
        p0[1].add_run(name_vl).bold = True
        pt_f1 = p0[1].paragraph_format  
        pt_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_f1.space_after = Pt(2)

        p0[2] = document.add_paragraph()
        p0[2].add_run('Безопасность производства работ на участках ВЛ '\
                    +'обеспечивается при следующих Схемах заземления ВЛ:').bold = True
        pt_f1 = p0[2].paragraph_format  
        pt_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_f1.space_after = Pt(2)

        p0[3] = document.add_paragraph()
        
        p0[3].add_run('СХЕМА '+NS+sv+' ЗАЗЕМЛЕНИЯ').bold = True
        pt_f1 = p0[3].paragraph_format  
        pt_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt_f1.space_after = Pt(2)

    # Информация о заземлении
    s1=[]
    s2=[]
    re_op_key = list(per_name.keys())    
    c1=False
    c2=False
    v=[]
    inf1 = []
    for h in range(len(inf)):
        
        ss1,ss2 =inf[h][1].split(Razd_simv,1)
        ss1=(ss1.strip()).replace(" ","\u00A0")
        ss2=(ss2.strip()).replace(" ","\u00A0")
        inf1.append(ss1+" "+Razd_simv+" "+ss2)

        if inf[h][0] ==1:
            if inf[h][5]:
                s1.append(ss1)
                c1=True
            else:
                s2.append(ss1)
                c2=True
                
            if inf[h][6]:
                s1.append(ss2)
                c1=True
            else:
                s2.append(ss2)
                c2=True
        elif inf[h][0] ==2:
            if inf[h][5]:
                s1.append(ss1)
                c1=True
            else:
                s2.append(ss1)
                c2=True
        elif inf[h][0] ==3:
            if inf[h][6]:
                s1.append(ss2)
                c1=True
            else:
                s2.append(ss2)
                c2=True
        #v.append(i)
        make_zy_list = []
        for j in range(len(lbsz[h])):
            if lbsz[h][j][0] in per_name[re_op_key[h]]:
                new_op = per_name[re_op_key[h]][lbsz[h][j][0]]
            else:
                new_op = lbsz[h][j][0]
            c1=True
            #s1.append('Опора №'+str(new_op)+' '+('('+lbsz[h][j][1]+' '+lbsz[h][j][2].replace(",",".")+')').replace(" ","\u00A0")+' на участке '+inf1[h])
            make_zy_list.append([new_op, ('('+lbsz[h][j][1]+' '+lbsz[h][j][2].replace(",",".")+')').replace(" ","\u00A0"),inf1[h]])


        vetvi_dict = {}
        for j in range(len(make_zy_list)):
            if make_zy_list[j][2] not in vetvi_dict:
                vetvi_dict[make_zy_list[j][2]] = [j]
            else:
                vetvi_dict[make_zy_list[j][2]].append(j)

        #if len(inf)==1:

        for key_name in vetvi_dict:
            for j in vetvi_dict[key_name]:
                s1.append('Опора №'+str(make_zy_list[j][0])+' '+make_zy_list[j][1])

            if len(inf)>1:
                s1[len(s1)-1]+= ' на участке '+key_name

    ign_f = set()
    for i in range(len(inf)):
        if grop:
            for j in range(len(f[i])):
                if inf[i][0]==1 and abs(inf[i][2]-inf[i][3])>=3:
                    if inf[i][2]==f[i][j][0][0]:
                        if f[i][j][0][1]==inf[i][2]+1 or f[i][j][0][1]==inf[i][2]-1:
                            ign_f.add((i,j))
                        elif inf[i][2]<inf[i][3]:
                            f[i][j][0][0]=inf[i][2]+1
                        elif inf[i][2]>inf[i][3]:
                            f[i][j][0][0]=inf[i][2]-1
                    if inf[i][3]==f[i][j][0][1]:
                        if f[i][j][0][0]==inf[i][3]-1 or f[i][j][0][0]==inf[i][3]+1:
                            ign_f.add((i,j))
                        elif inf[i][2]<inf[i][3]:
                            f[i][j][0][1]=inf[i][3]-1
                        elif inf[i][2]>inf[i][3]:
                            f[i][j][0][1]=inf[i][3]+1
                if inf[i][0]==2 and abs(inf[i][2]-inf[i][3])>=2:
                    if inf[i][2]==f[i][j][0][0]:
                        if f[i][j][0][1]==inf[i][2]+1 or f[i][j][0][1]==inf[i][2]-1:
                            f[i][j] = None
                        elif inf[i][2]<inf[i][3]:
                            f[i][j][0][0]=inf[i][2]+1
                        elif inf[i][2]>inf[i][3]:
                            f[i][j][0][0]=inf[i][2]-1
                if inf[i][0]==3 and abs(inf[i][2]-inf[i][3])>=2:
                    if inf[i][3]==f[i][j][0][1]:
                        if f[i][j][0][0]==inf[i][3]-1 or f[i][j][0][0]==inf[i][3]+1:
                            f[i][j] = None
                        elif inf[i][2]<inf[i][3]:
                            f[i][j][0][1]=inf[i][3]-1
                        elif inf[i][2]>inf[i][3]:
                            f[i][j][0][1]=inf[i][3]+1

    # Цикл по ветвям
    p1 = []
    if tb1:
        for i in range(len(inf)):
            p1.append(['','','','',''])

            # Имя ветви
            p1[i][0]=document.add_paragraph()
            p1[i][0].add_run(inf1[i]).bold = True
            pt_f = p1[i][0].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_f.space_after = Pt(2)


            

            if c1:
                p1[i][1]=document.add_paragraph()
                pt_f = p1[i][1].paragraph_format
                pt_f.first_line_indent = Mm(12.5)
                pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Выравнивание по ширине
                p1[i][1].add_run('ВЛ заземлена: ').bold = True
                for j in range(len(s1)):
                    if j == len(s1)-1:
                        p1[i][1].add_run(s1[j]+'.')
                    else:
                        p1[i][1].add_run(s1[j]+', ')
            if c2:
                p1[i][2]=document.add_paragraph()
                pt_f = p1[i][2].paragraph_format
                pt_f.first_line_indent = Mm(12.5)
                pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Выравнивание по ширине
                p1[i][2].add_run('ВЛ разземлена: ').bold = True
                for j in range(len(s2)):
                    if j == len(s2)-1:
                        p1[i][2].add_run(s2[j]+'.')
                    else:
                        p1[i][2].add_run(s2[j]+', ')

            # Картинка схемы заземления
            p1[i][3]=document.add_picture('result_schemes/'+str(i)+'.jpg', width=Inches(6.5))
            #last_paragraph = document.paragraphs[-1] 
            #last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Создаём шапку таблицы
            p1[i][4]=document.add_table(rows=1, cols=2, style='Table Grid') # Таблица с размерами и стилем
            
            p_t=p1[i][4].style.paragraph_format
            p_t.keep_with_next = True # Если невмещается таблица на текущую страницу, переносит её на ноыую
            
            row_1=p1[i][4].rows[0] # Обращаемся к строке заголовка

            row_1.cells[0].text = 'Участки ВЛ' # Присваиваем текст
            r = row_1.cells[0].paragraphs[0].runs[0] # Делаем его жирным
            row_1.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_vertical_alignment(row_1.cells[0], align="center") # Выравнивание содержимого по вертикали
            r.font.bold =True

            row_1.cells[1].text = 'Разрешение на выполнение работ'
            row_1.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_vertical_alignment(row_1.cells[1], align="center") # Выравнивание содержимого по вертикали
            r = row_1.cells[1].paragraphs[0].runs[0]
            r.font.bold =True

            ss1,ss2 =inf[i][1].split(Razd_simv,1)
            ss1=(ss1.strip()).replace(" ","\u00A0")
            ss2=(ss2.strip()).replace(" ","\u00A0")

            
            # ПС слева (если есть)
            if inf[i][0] ==1 or inf[i][0] ==2:
                row_cells = p1[i][4].add_row().cells # Добавляем в таблицу новую строку
                row_cells[0].text = ss1
                set_cell_vertical_alignment(row_cells[0], align="center") # Выравнивание содержимого по вертикали
                if f2[i][inf[i][2]] or (zpps==True and inf[i][5]==False and (inf[h][0]==1 or inf[h][0]==2)):
                    row_cells[1].text = 'Запрещено'
                    r = row_cells[1].paragraphs[0].runs[0]
                    r.font.bold =True
                    set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали
                else:
                    row_cells[1].text = 'Разрешено'
                    set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали

            # Участки ВЛ
            
            
            for j in range(len(f[i])):   
                if (i,j) in ign_f: continue 

                row_cells = p1[i][4].add_row().cells # Добавляем в таблицу новую строку
                
                if f[i][j][0][0] in per_name[re_op_key[i]]:
                    new_op1 = per_name[re_op_key[i]][f[i][j][0][0]]
                else:
                    new_op1 = f[i][j][0][0]
                if f[i][j][0][1] in per_name[re_op_key[i]]:
                    new_op2 = per_name[re_op_key[i]][f[i][j][0][1]]
                else:
                    new_op2 = f[i][j][0][1]
                row_cells[0].text = 'Опоры № '+str(new_op1)+'-'+str(new_op2)
                set_cell_vertical_alignment(row_cells[0], align="center") # Выравнивание содержимого по вертикали
                if f[i][j][1]:
                    row_cells[1].text = 'Запрещено'
                    set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали
                    r = row_cells[1].paragraphs[0].runs[0]
                    r.font.bold =True
                else:
                    row_cells[1].text = 'Разрешено'
                    set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали
            # ПС справа (если есть)
            if inf[i][0] ==1 or inf[i][0] ==3:
                row_cells = p1[i][4].add_row().cells # Добавляем в таблицу новую строку
                row_cells[0].text = ss2
                set_cell_vertical_alignment(row_cells[0], align="center") # Выравнивание содержимого по вертикали
                if f2[i][inf[i][3]] or (zpps==True and inf[i][6]==False and (inf[h][0]==1 or inf[h][0]==3)):
                    row_cells[1].text = 'Запрещено'
                    set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали
                    r = row_cells[1].paragraphs[0].runs[0]
                    r.font.bold =True
                else:
                    row_cells[1].text = 'Разрешено'
                    set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали
            document.add_paragraph()

        if tb2: document.add_page_break() # Разрыв страницы

    # Эпюры
    if tb2:
        # Информация о заземлении
        s1=[]
        s3=[]
        s2=[]
            
        c1=False
        c2=False
        c3=False

        for i in range(len(inf)):
            ss1,ss2 =inf[i][1].split(Razd_simv,1)
            ss1=(ss1.strip()).replace(" ","\u00A0")
            ss2=(ss2.strip()).replace(" ","\u00A0")


            
            if inf[i][0] ==1:
                if inf[i][5]:
                    s1.append(ss1)
                    c1=True
                else:
                    s2.append(ss1)
                    c2=True
                    
                if inf[i][6]:
                    s1.append(ss2)
                    c1=True
                else:
                    s2.append(ss2)
                    c2=True
            elif inf[i][0] ==2:
                if inf[i][5]:
                    s1.append(ss1)
                    c1=True
                else:
                    s2.append(ss1)
                    c2=True
            elif inf[i][0] ==3:
                if inf[i][6]:
                    s1.append(ss2)
                    c1=True
                else:
                    s2.append(ss2)
                    c2=True


            make_zy_list = []

            for j in range(len(lbsz[i])):
                if lbsz[i][j][0] in per_name[re_op_key[i]]:
                    new_op = per_name[re_op_key[i]][lbsz[i][j][0]]
                else:
                    new_op = lbsz[i][j][0]
                c3=True
                #s3.append('№'+str(new_op)+' '+('('+lbsz[i][j][1]+' '+lbsz[i][j][2].replace(",",".")+')').replace(" ","\u00A0")+' на участке '+inf1[i])
                
                make_zy_list.append([new_op, ('('+lbsz[i][j][1]+' '+lbsz[i][j][2].replace(",",".")+')').replace(" ","\u00A0"),inf1[i]])

            vetvi_dict = {}
            for j in range(len(make_zy_list)):
                if make_zy_list[j][2] not in vetvi_dict:
                    vetvi_dict[make_zy_list[j][2]] = [j]
                else:
                    vetvi_dict[make_zy_list[j][2]].append(j)

            #if len(inf)==1:

            for key_name in vetvi_dict:
                for j in vetvi_dict[key_name]:
                    s3.append('№'+str(make_zy_list[j][0])+' '+make_zy_list[j][1])

                if len(inf)>1:
                    s3[len(s3)-1]+= ' на участке '+key_name

        
        # Цикл по ветвям
        p = []
        for i in range(len(inf)):
            p.append(['','','','','','','','','','','','','','','','','',''])

            # Имя ВЛ
            p[i][15]=document.add_paragraph()
            p[i][15].add_run(name_vl).bold = True
            pt_f = p[i][15].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_f.space_after = Pt(2)
            
            # Заголовок
            p[i][0]=document.add_paragraph()

            if len(inf)==1:
                p[i][0].add_run('СХЕМА '+NS+sv+' ЗАЗЕМЛЕНИЯ').bold = True 
            else:
                p[i][0].add_run('СХЕМА '+NS+sv+' ЗАЗЕМЛЕНИЯ').bold = True
                #p[i][0].add_run('СХЕМА '+NS+sv+'.'+str(i+1)+' ЗАЗЕМЛЕНИЯ').bold = True
            
            pt_f = p[i][0].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_f.space_after = Pt(2)

            # Рабочий участок
            p[i][1]=document.add_paragraph()
            p[i][1].add_run('РАБОЧИЙ УЧАСТОК').bold = True
            pt_f = p[i][1].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_f.space_after = Pt(2)

            # ПС и опоры на которых можно работать
            p[i][2]=document.add_paragraph('Опоры № ')

            ss1,ss2 =inf[i][1].split(Razd_simv,1)
            ss1=(ss1.strip()).replace(" ","\u00A0")
            ss2=(ss2.strip()).replace(" ","\u00A0")

            
            s4=[]
            s5=[]
            s6=[]
            s7=[]
            c4=False
            c5=False
            c6=False
            c7=False
            
            # Участки ВЛ
            for j in range(len(f[i])):
                if (i,j) in ign_f: continue 
                if f[i][j][0][0] in per_name[re_op_key[i]]:
                    new_op1 = per_name[re_op_key[i]][f[i][j][0][0]]
                else:
                    new_op1 = f[i][j][0][0]
                if f[i][j][0][1] in per_name[re_op_key[i]]:
                    new_op2 = per_name[re_op_key[i]][f[i][j][0][1]]
                else:
                    new_op2 = f[i][j][0][1]
                    
                if not f[i][j][1]:
                    s4.append(str(new_op1)+'-'+str(new_op2))
                    c4 = True
                elif f[i][j][1]:
                    s6.append(str(new_op1)+'-'+str(new_op2))
                    c6 = True
            
                  
            # ПС слева (если есть)
            if inf[i][0] ==1 or inf[i][0] ==2:
                if f2[i][inf[i][2]] or (zpps==True and inf[i][5]==False and (inf[i][0]==1 or inf[i][0]==2)):
                    s7.append(ss1)
                    c7=True
                else:
                    s5.append(ss1)
                    c5=True
                
                    
            # ПС справа (если есть)
            if inf[i][0] ==1 or inf[i][0] ==3:
                if f2[i][inf[i][3]] or (zpps==True and inf[i][6]==False and (inf[i][0]==1 or inf[i][0]==3)):
                    s7.append(ss2)
                    c7=True
                else:
                    s5.append(ss2)
                    c5=True
                
                    
            if c4:
                for j in range(len(s4)):
                    if j == len(s4)-1 and c5:
                        p[i][2].add_run(s4[j]+', ')
                    elif j == len(s4)-1 and not c5:
                        p[i][2].add_run(s4[j])
                    else:
                        p[i][2].add_run(s4[j]+', ')

            if c5:
                for j in range(len(s5)):
                    if j == len(s5)-1:
                        p[i][2].add_run(s5[j])
                    else:
                        p[i][2].add_run(s5[j]+', ')

            pt_f = p[i][2].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_f.space_after = Pt(2)
            
            # Схема заземления вл
            p[i][3]=document.add_paragraph()
            p[i][3].add_run('СХЕМА ЗАЗЕМЛЕНИЯ ВЛ').bold = True
            pt_f = p[i][3].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Картинка схемы заземления
            p[i][4]=document.add_picture('result_schemes/'+str(i)+'.jpg', width=Inches(6.5))
            #last_paragraph = document.paragraphs[-1] 
            #last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Карта заземления
            p[i][5]=document.add_paragraph()
            p[i][5].add_run('КАРТА ЗАЗЕМЛЕНИЯ').bold = True
            pt_f = p[i][5].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_f.space_after = Pt(2)

            


            k=1
            # Пункт 1 (Инфа о заземлении)
            p[i][6]=document.add_paragraph('1. ')
            pt_f = p[i][6].paragraph_format
            pt_f.first_line_indent = Mm(12.5)
            pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Выравнивание по ширине
            if c1:
                p[i][6].add_run('Заземлить ВЛ на ПЗ на ')
                for j in range(len(s1)):
                    if j == len(s1)-1:
                        p[i][6].add_run(s1[j]+'. ')
                    else:
                        p[i][6].add_run(s1[j]+', ') 
                k =2
            if c3:
                p[i][6].add_run('Заземлить ВЛ на опор%s ' % ('e' if len(s3)==1 else 'ах') )
                for j in range(len(s3)):
                    if j == len(s3)-1:
                        p[i][6].add_run(s3[j]+'. ')
                    else:
                        p[i][6].add_run(s3[j]+', ')
                k =2
            if c2:
                p[i][6].add_run('Не заземлять ВЛ на ')
                for j in range(len(s2)):
                    if j == len(s2)-1:
                        p[i][6].add_run(s2[j]+'. ')
                    else:
                        p[i][6].add_run(s2[j]+', ')
                k =2

            # Пункт 2 (Заземление ЛЗ на рабочих участках)
            if c4:
                p[i][7]=document.add_paragraph(str(k)+'. Заземление каждого рабочего места на ВЛ - заземление типа ЛЗ.')
                k+=1
                pt_f = p[i][7].paragraph_format
                pt_f.first_line_indent = Mm(12.5)
                pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Выравнивание по ширине
            # Пункт 3 (Разрешение работ)
            if (c4 and c5) or (c4 or c5):
                p[i][8]=document.add_paragraph(str(k)+'. ')
                pt_f = p[i][8].paragraph_format
                pt_f.first_line_indent = Mm(12.5)
                pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Выравнивание по ширине
                if c4:
                    if i == 0:
                        p[i][8].add_run('Разрешаются работы на участках ВЛ опоры № ')
                    else:
                        p[i][8].add_run('Разрешаются работы на участках ответвления от ВЛ опоры № ')
                    for j in range(len(s4)):
                        if j == len(s4)-1:
                            p[i][8].add_run(s4[j]+'. ')
                        else:
                            p[i][8].add_run(s4[j]+', ')
                if c5:
                    p[i][8].add_run('Разрешаются работы на линейном оборудовании ')
                    for j in range(len(s5)):
                        if j == len(s5)-1:
                            p[i][8].add_run(s5[j]+'. ')
                        else:
                            p[i][8].add_run(s5[j]+', ')
                    p[i][8].add_run('При работах на линейных разъединителях дополнительно устанавливается заземление типа ДЗ. ')
                k +=1

            # Пункт 4 (если есть безопасные участки вл)
            if c4:
                p[i][9]=document.add_paragraph(str(k)+'. Разрешается работа на участке ВЛ до 2 км'\
                                            +' с установкой заземления типа ЛЗ с двух сторон участка,'\
                                            +' при условии, что их установка производится в пределах рабочего участка ВЛ.')
                k+=1
                pt_f = p[i][9].paragraph_format
                pt_f.first_line_indent = Mm(12.5)
                pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Выравнивание по ширине
            # Пункт 5 (если есть запрещённые участки)
            if (c6 and c7) or (c6 or c7):
                p[i][10]=document.add_paragraph(str(k)+'. ')
                pt_f = p[i][10].paragraph_format
                pt_f.first_line_indent = Mm(12.5)
                pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Выравнивание по ширине
                if c6:
                    if i == 0:
                        p[i][10].add_run('Запрещаются работы на участках ВЛ опоры № ')
                    else:
                        p[i][10].add_run('Запрещаются работы на участках ответвления от ВЛ опоры № ')
                    for j in range(len(s6)):
                        if j == len(s6)-1:
                            p[i][10].add_run(s6[j]+'. ')
                        else:
                            p[i][10].add_run(s6[j]+', ')
                if c7:
                    p[i][10].add_run('Запрещаются работы на линейном оборудовании ')
                    for j in range(len(s7)):
                        if j == len(s7)-1:
                            p[i][10].add_run(s7[j]+'. ')
                        else:
                            p[i][10].add_run(s7[j]+', ')
                k +=1

            # Пункт 6 (при установке ЛЗ 25 В если есть безопасные участки)
            #if c4:
                #p[i][11]=document.add_paragraph(str(k)+'. При установке заземления типа ЛЗ напряжение на рабочих '\
                                                #+'местах не превысит нормируемой величины 25 В.')

            # Распеределение напряжения по длинне ВЛ
            p[i][12]=document.add_paragraph()
            p[i][12].add_run('СХЕМА СБЛИЖЕНИЯ').bold = True
            pt_f = p[i][12].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_f.space_after = Pt(2)
            pt_f.space_before = Pt(10)

            
            # Картинка эпюры напряжения
            im = Image.open('gr_sb/'+str(i)+'.jpg')
            im.crop((80,0,1610,500)).save('gr_sb/obr'+str(i)+'.jpg')
            p[i][13]=document.add_picture('gr_sb/obr'+str(i)+'.jpg', width=Inches(6.25))
            last_paragraph = document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


            p[i][14]=document.add_paragraph()
            p[i][14].add_run('РАСПРЕДЕЛЕНИЕ НАПРЯЖЕНИЯ ПО ДЛИНЕ ВЛ').bold = True
            pt_f = p[i][14].paragraph_format
            pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt_f.space_after = Pt(2)
            pt_f.space_before = Pt(10)

            
            # Картинка эпюры напряжения
            im = Image.open('images_grafik/'+str(i)+'.jpg')
            im.crop((80,15,1610,500)).save('images_grafik/obr'+str(i)+'.jpg', "PNG")
            p[i][15]=document.add_picture('images_grafik/obr'+str(i)+'.jpg', width=Inches(6.25))
            last_paragraph = document.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Разрыв страницы если нужно
            if i != len(inf)-1:
                document.add_page_break() # Разрыв страницы
            
    # Сохраняем документ
    if Trig:
        document.add_page_break() # Разрыв страницы
        return document
    else:
        document.save(fname)
        return None
    
    
