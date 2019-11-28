# pylint: disable=E0611
# pylint: disable=E1101
from docx import Document
from docx.shared import Pt, Mm
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE
from docx.oxml.shared import OxmlElement, qn
from PIL import Image


def WordOid(vvl, sp_pzy,nm_d, ivl_gr, fname,nVL,grop, D_file, Trig):

    # Создаём объект документа
    if Trig:
        document = D_file
    else:
        document = Document()

    # Наследуем стиль и изменяем его
    style = document.styles['Normal'] # Берём стиль Нормальный


    f0 = style.font # Переменная для изменения параметров стиля
    f0.name = 'Arial' # Шрифт
    f0.size = Pt(12) # Размер шрифта

    pf = style.paragraph_format
    pf.line_spacing = Pt(0) # Междустрочный интервал
    pf.space_after = Pt(0) # Интервал после абзаца

    # Задаём свойство полей
    sections = document.sections
    s = sections[0]
    s.left_margin = Mm(30)
    s.right_margin = Mm(15)
    s.top_margin = Mm(20)
    s.bottom_margin = Mm(20)
    s.page_height = Mm(297)
    s.page_width = Mm(210)

    
    p = ['','','','','','','','','']

    # Картинка схемы сближения
    p[0]=document.add_picture('Pictures/sh2.jpg', height=Inches(9.8)) #width=Inches(7)
    document.add_page_break() # Разрыв страницы
    
    # Заголовок
    p[1]=['','']
    p[1][0]=document.add_paragraph()
    p[1][0].add_run('Основные исходные данные для расчёта наведенного напряжения на ').bold = True
    pt_f = p[1][0].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p[1][1]=document.add_paragraph()
    p[1][1].add_run(nVL).bold = True
    pt_f = p[1][1].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_f.space_after = Pt(10)
    
    """ 
    # Схема сближения
    p[2]=document.add_paragraph()
    p[2].add_run('Схема сближения').bold = True
    pt_f = p[2].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_f.space_after = Pt(10) """

    
    """
    # Условные обозначения
    p[3]=document.add_paragraph('На схеме сближения приняты следующие условные обозначения:')
    pt_f = p[3].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
    
    # Список условных обозначений
    p[4]=[[],[]]
    for i in range(len(sp_pzy)):
        p[4][0].append(document.add_paragraph(sp_pzy[i][0]+' - '+sp_pzy[i][1]))
        pt_f = p[4][0][i].paragraph_format
        pt_f.first_line_indent = Mm(20)
        pt_f.alignment = WD_ALIGN_PARAGRAPH.LEFT        

    i = -1
    for key in nm_d:
        i+=1
        p[4][1].append(document.add_paragraph(key+' - '+nm_d[key]))
        pt_f = p[4][1][i].paragraph_format
        pt_f.first_line_indent = Mm(20)
        pt_f.alignment = WD_ALIGN_PARAGRAPH.LEFT
    """
    nm_d_k = list(nm_d.keys())
    # Загаловок таблицы Мощностей
    
    p[5]=document.add_paragraph()
    p[5].add_run('Мощности влияющих ВЛ').bold = True
    pt_f = p[5].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_f.space_after = Pt(10)
    pt_f.space_before = Pt(10)    

    # Создаём шапку таблицы
    p[6]=document.add_table(rows=2, cols=5, style='Table Grid')

    # Функция выравнивания содержимого таблицы по вертикали
    def set_cell_vertical_alignment(cell, align="center"): 
        try:   
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcValign = OxmlElement('w:vAlign')  
            tcValign.set(qn('w:val'), align)  
            tcPr.append(tcValign)
            return True 
        except:
            #traceback.print_exc()             
            return False
    
    p[6].cell(0, 0).merge(p[6].cell(1, 0)) # Объеденяем ячейки по вертикали(Условное обозначение)
    p[6].cell(0, 1).merge(p[6].cell(1, 1)) # Объеденяем ячейки по вертикали(Наименование ВЛ)
    p[6].cell(0, 2).merge(p[6].cell(1, 2)) # Объеденяем ячейки по вертикали (Напряжение ВЛ)
    p[6].cell(0, 3).merge(p[6].cell(0, 4)) # Объеденяем ячейки по горизонтали (Мощность)

    p[6].cell(0, 0).text = 'L №'
    p[6].cell(0, 0).width = Inches(0.2) # Задаёт ширину ячейк
    p[6].cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Делает текст по центру ширины ячейки
    set_cell_vertical_alignment(p[6].cell(0, 0), align="center") # Выравнивание содержимого по вертикали

    p[6].cell(0, 1).text = 'Наименование ВЛ'
    p[6].cell(0, 1).width = Inches(8) # Задаёт ширину ячейк
    p[6].cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Делает текст по центру ширины ячейки
    set_cell_vertical_alignment(p[6].cell(0, 1), align="center") # Выравнивание содержимого по вертикали
  
    p[6].cell(0, 2).text = 'Напряжение ВЛ, кВ'
    p[6].cell(0, 2).width = Inches(1)
    p[6].cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(p[6].cell(0, 2), align="center") # Выравнивание содержимого по вертикали
    
    p[6].cell(0, 3).text = 'Мощность'
    p[6].cell(0, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p[6].cell(1, 3).width = Inches(2)
    set_cell_vertical_alignment(p[6].cell(0, 3), align="center") # Выравнивание содержимого по вертикали
    
    p[6].cell(1, 3).text = 'Активная, МВт'
    p[6].cell(1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p[6].cell(1, 3).width = Inches(1)
    set_cell_vertical_alignment(p[6].cell(1, 3), align="center") # Выравнивание содержимого по вертикали
    
    p[6].cell(1, 4).text = 'Реактивная, МВар'
    p[6].cell(1, 4).width = Inches(1)
    p[6].cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(p[6].cell(1, 4), align="center") # Выравнивание содержимого по вертикали

    
    # Заполняем  таблицу
    j = -1
    for i in range(len(vvl)):
        if vvl[i][0] != 'Конфигурация опоры':
            j+=1
            row_cells = p[6].add_row().cells

            row_cells[0].text = nm_d_k[j]
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_vertical_alignment(row_cells[0], align="center") # Выравнивание содержимого по вертикали
            
            row_cells[1].text = vvl[i][0]
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали

            row_cells[2].text = vvl[i][3].replace(",",".")
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_vertical_alignment(row_cells[2], align="center") # Выравнивание содержимого по вертикали

            row_cells[3].text = vvl[i][4].replace(",",".")
            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_vertical_alignment(row_cells[3], align="center") # Выравнивание содержимого по вертикали

            row_cells[4].text = vvl[i][5].replace(",",".")
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_vertical_alignment(row_cells[4], align="center") # Выравнивание содержимого по вертикали

    document.add_paragraph()
    
    # Загаловок таблицы Мощностей
    p[7]=document.add_paragraph()
    p[7].add_run('Сопротивление заземляющих устройств подстанций').bold = True
    pt_f = p[7].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_f.space_after = Pt(10)

    # Таблица ЗУ
    p[8]=document.add_table(rows=1, cols=3, style='Table Grid')

    p[8].cell(0, 0).text = 'ПС №'
    p[8].cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(p[8].cell(0, 0), align="center") # Выравнивание содержимого по вертикали
    

    p[8].cell(0, 1).text = 'Наименование ПС'
    p[8].cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p[8].cell(0, 1).width = Inches(10)
    set_cell_vertical_alignment(p[8].cell(0, 1), align="center") # Выравнивание содержимого по вертикали
    

    p[8].cell(0, 2).text = 'Сопротивление ЗУ, Ом'
    p[8].cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p[8].cell(0, 2).width = Inches(3)
    set_cell_vertical_alignment(p[8].cell(0, 2), align="center") # Выравнивание содержимого по вертикали

    
    for i in range(len(sp_pzy)):
        row_cells = p[8].add_row().cells
        
        row_cells[0].text = sp_pzy[i][0]
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(row_cells[0], align="center") # Выравнивание содержимого по вертикали
                        
        row_cells[1].text = sp_pzy[i][1].replace(" ","\u00A0") # Заменяем обычные пробелы на неразрывные
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали

        row_cells[2].text = sp_pzy[i][2].replace(",",".")
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(row_cells[2], align="center") # Выравнивание содержимого по вертикали

                            
    
    #document.add_page_break() # Разрыв страницы
        
    # Сохраняем документ
    if Trig:
        document.add_page_break() # Разрыв страницы
        return document
    else:
        document.save(fname)
        return None
