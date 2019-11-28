from docx import Document
from docx.shared import Pt, Mm
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE
from docx.oxml.shared import OxmlElement, qn
from PIL import Image





def WordSxT(sp_t,sp_op,d_PZ,ku1,fname,name_vl, D_file, Trig):

    # Создаём объект документа
    if Trig: 
        document = D_file
    else:
        document = Document()

    # Наследуем стиль и изменяем его
    style = document.styles['Normal'] # Берём стиль Нормальный

    f0 = style.font # Переменная для изменения параметров стиля
    f0.name = 'Arial' # Шрифт
    f0.size = Pt(12) # Размер шрифта

    pf = style.paragraph_format
    pf.line_spacing = Pt(0) # Междустрочный интервал
    pf.space_after = Pt(0) # Интервал после абзаца

    # Задаём свойство полей
    sections = document.sections
    s = sections[0]
    s.left_margin = Mm(30)
    s.right_margin = Mm(15)
    s.top_margin = Mm(20)
    s.bottom_margin = Mm(20)
    s.page_height = Mm(297)
    s.page_width = Mm(210)

    
    p = ['','','','','','','','','','','','','','','',[]]

    # Название вл
    p[0]=document.add_paragraph()
    p[0].add_run(name_vl).bold = True
    pt_f = p[0].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_f.space_after = Pt(2)

    # Заголовок
    p[1]=document.add_paragraph()
    p[1].add_run('СХЕМА №3 ЗАЗЕМЛЕНИЯ ВЛ').bold = True
    pt_f = p[1].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_f.space_after = Pt(10)

    # Введение
    p[2]=document.add_paragraph('В таблице информация о безoпасных для работы участках при установке БЗ на ')
    p[2].add_run('перечисленных опорах. Следуйте указаниям, приведённым в карте заземления ВЛ.')
    pt_f = p[2].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pt_f.space_after = Pt(10)
    

    p[3]=document.add_paragraph()
    p[3].add_run('КАРТА ЗАЗЕМЛЕНИЯ ВЛ').bold = True
    pt_f = p[3].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt_f.space_after = Pt(10)
    
    p[4]=document.add_paragraph('1. В таблице в колонке "БЗ на опоре №" выбрать опору, на которую необходимо ')
    p[4].add_run('установить БЗ. Последующие пункты обращаются к строке таблицы, в которой находится ')
    p[4].add_run('выбранная опора. Место установки заземлителя типа БЗ на рабочем участке выбирается ')
    p[4].add_run('исходя из удобства проведения работ.')
    pt_f = p[4].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p[5]=document.add_paragraph('2.1. Заземлить ВЛ на ПЗ на ПС указанной в колонке "ПС". ')
    p[5].add_run('На прочих ПС ВЛ - ')
    p[5].add_run('разземлена.').bold = True
    pt_f = p[5].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p[6]=document.add_paragraph('2.2. Заземлить ВЛ на заземлитель типа БЗ на выбранной опоре. ')
    pt_f = p[6].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p[7]=document.add_paragraph('2.3. После заземления ВЛ на БЗ, разземлить ВЛ на ПС указанной ')
    p[7].add_run('в колонке "ПС".')
    pt_f = p[7].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p[8]=document.add_paragraph('3. Заземление каждого рабочего места на ВЛ - заземление типа ЛЗ.')
    pt_f = p[8].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    p[9]=document.add_paragraph('4. При установке БЗ на выбранной опоре разрешается выполнение работ ')
    p[9].add_run('только ').bold = True
    p[9].add_run('на участках указанных в колонке "Рабочий участок". ')
    p[9].add_run('Если безопасных участков не будет, то в колонке будет указано - ')
    p[9].add_run('запрещено.').bold = True
    pt_f = p[9].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p[10]=document.add_paragraph('5.1. Установка ЛЗ разрешается только в пределах участков, указанных в ')
    p[10].add_run('колонке "Рабочий участок".')
    pt_f = p[10].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p[11]=document.add_paragraph('5.2. Разрешается работа на участке ВЛ до 2 км с установкой заземления ')
    p[11].add_run('типа ЛЗ с двух сторон участка, если это не противоречит пункту 5.1. ')
    p[11].add_run('В противном случае необходимо ограничиться участком указанном в колонке ')
    p[11].add_run('"Рабочий участок".')
    pt_f = p[11].paragraph_format
    pt_f.first_line_indent = Mm(12.5)
    pt_f.space_after = Pt(10)
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
 
    p[14]=document.add_paragraph('Для названий в таблице приняты следующие условные обозначения: ')
    pt_f = p[14].paragraph_format
    pt_f.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY # Выравнивание по ширине
    pt_f.first_line_indent = Mm(12.5)
    for key in d_PZ:
        pp = document.add_paragraph(d_PZ[key]+' - '+key.replace(" ","\u00A0"))
        pt_f = pp.paragraph_format
        pt_f.first_line_indent = Mm(20)
        p[15].append(pp)
    """
    for key in ku1:
        pp = document.add_paragraph(ku1[key]+' - '+key) 
        pt_f = pp.paragraph_format
        pt_f.first_line_indent = Mm(20)
        p[15].append(pp)
    """
    nm_line = inv_map = {v: k for k, v in ku1.items()}
    
 

    document.add_page_break() # Разрыв страницы

    #p[12]=document.add_paragraph('Таблица безопасных для работы участков по Схеме №3')
    #document.add_paragraph('')
    # Создаём шапку таблицы
    p[13]=document.add_table(rows=2, cols=6, style='Table Grid')

    # Функция выравнивания содержимого таблицы по вертикали
    def set_cell_vertical_alignment(cell, align="center"): 
        try:   
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcValign = OxmlElement('w:vAlign')  
            tcValign.set(qn('w:val'), align)  
            tcPr.append(tcValign)
            return True 
        except:
            #traceback.print_exc()             
            return False
    
    p[13].cell(0, 0).merge(p[13].cell(0, 5))
    p[13].cell(0, 0).text = 'Таблица безопасных для работы участков по Схеме №3'
    set_cell_vertical_alignment(p[13].cell(0, 0), align="center")
    p[13].cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p[13].cell(1, 0).text = 'БЗ на опоре №'
    #p[6].cell(0, 0).width = Inches(6) # Задаёт ширину ячейк
    set_cell_vertical_alignment(p[13].cell(1, 0), align="center")
    p[13].cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Делает текст по центру ширины ячейки
  
    p[13].cell(1, 1).text = 'ПС'
    set_cell_vertical_alignment(p[13].cell(1, 1), align="center")
    #p[12].cell(0, 1).width = Inches(1)
    p[13].cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
    
    p[13].cell(1, 2).text = 'Рабочий участок'
    set_cell_vertical_alignment(p[13].cell(1, 1), align="center")
    p[13].cell(1, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #p[12].cell(1, 2).width = Inches(2)
    p[13].cell(1, 3).text = 'БЗ на опоре №'
    #p[6].cell(0, 0).width = Inches(6) # Задаёт ширину ячейк
    set_cell_vertical_alignment(p[13].cell(1, 3), align="center")
    p[13].cell(1, 3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Делает текст по центру ширины ячейки
  
    p[13].cell(1, 4).text = 'ПС'
    set_cell_vertical_alignment(p[13].cell(1, 4), align="center")
    #p[12].cell(0, 1).width = Inches(1)
    p[13].cell(1, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
    
    p[13].cell(1, 5).text = 'Рабочий участок'
    set_cell_vertical_alignment(p[13].cell(1, 5), align="center")
    p[13].cell(1, 5).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    #p[12].cell(1, 2).width = Inches(2)



    
    # Заполняем  таблицу
    for i in range(len(sp_op)):
        
        if len(sp_op[i])==1:
            j = False
            row_cells = p[13].add_row().cells
            
            row_cells[0].merge(row_cells[5])
            row_cells[0].text = nm_line[sp_op[i][0]]
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_vertical_alignment(row_cells[0], align="center") # Выравнивание содержимого по вертикали
        else:
            j = True if not j else False

            if j:
                row_cells = p[13].add_row().cells
                
                row_cells[0].text = sp_op[i][0]
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(row_cells[0], align="center") # Выравнивание содержимого по вертикали
                r = row_cells[0].paragraphs[0].runs[0]
        
                row_cells[1].text = sp_op[i][1]
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(row_cells[1], align="center") # Выравнивание содержимого по вертикали


                row_cells[2].text = sp_op[i][2]
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(row_cells[2], align="center") # Выравнивание содержимого по вертикали
                if sp_op[i][2]=='Запрещено':
                    r = row_cells[2].paragraphs[0].runs[0]
                    r.font.bold =True

            else:
                row_cells[3].text = sp_op[i][0]
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(row_cells[3], align="center") # Выравнивание содержимого по вертикали
                r = row_cells[3].paragraphs[0].runs[0]
        
                row_cells[4].text = sp_op[i][1]
                row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(row_cells[4], align="center") # Выравнивание содержимого по вертикали


                row_cells[5].text = sp_op[i][2]
                row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(row_cells[5], align="center") # Выравнивание содержимого по вертикали
                if sp_op[i][2]=='Запрещено':
                    r = row_cells[5].paragraphs[0].runs[0]
                    r.font.bold =True
                    


    # Сохраняем документ
    if Trig:
        document.add_page_break() # Разрыв страницы
        return document
    else:
        document.save(fname)
        return None
